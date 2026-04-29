#!/usr/bin/env python
"""
题库导入脚本
用于从 Word 文档中解析题目并导入到数据库
"""
import os
import sys
import django
from pathlib import Path

# 设置 Django 环境
sys.path.insert(0, str(Path(__file__).parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from docx import Document
from apps.questions.models import Question, QuestionCategory


def parse_docx_file(file_path):
    """解析 docx 文件，提取题目内容"""
    print(f"\n正在解析文件: {file_path}")
    doc = Document(file_path)

    print(f"文档段落数: {len(doc.paragraphs)}")
    print("\n前 20 个段落内容:")
    print("-" * 80)

    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"{i+1}. {para.text[:100]}")

    print("-" * 80)


def main():
    """主函数"""
    # 题库文件夹路径
    question_dir = Path(__file__).parent.parent / "题库"

    if not question_dir.exists():
        print(f"错误: 题库文件夹不存在: {question_dir}")
        return

    # 获取所有 docx 文件
    docx_files = list(question_dir.glob("*.docx"))

    print(f"找到 {len(docx_files)} 个 docx 文件")

    if docx_files:
        # 先查看第一个文件的内容
        parse_docx_file(docx_files[0])


if __name__ == "__main__":
    main()
